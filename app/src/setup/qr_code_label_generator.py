from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
import json
import pathlib
import math
import tempfile

from src.setup.qr_code import QrCodeImageGenerator


class QrCodeLabelGenerator:

    def __init__(self) -> None:
        self.qr_generator = QrCodeImageGenerator()
        self.qr_code_path = tempfile.TemporaryDirectory()

    def create_document(self, label_data, output_path, url):
        document = Document()

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(8)

        n_cols = 3
        n_rows = math.ceil(len(label_data.keys())/n_cols)

        self._set_document_margins(document, 25, 25, 25, 25)

        # Create a table with the desired dimensions
        table = document.add_table(rows=n_rows, cols=n_cols)
        table.style = 'TableGrid'
        table.autofit = False  # Disable autofit

        # Set cell dimensions to 50x50mm
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Mm(50)
            for cell in row.cells:
                cell.width = Mm(50)
    
        for (i, data) in enumerate(label_data):

            row_index = i // n_cols
            col_index = i % n_cols

            cell = table.cell(row_index, col_index)
            cell.paragraphs[0].style = document.styles['Normal']
            cell.paragraphs[0].alignment = 1 

            width= Mm(40)
            height= Mm(40)

            token = label_data[data]["token"]

            text = f'{url}/\nregistration?token={token}'
            image_file_path = pathlib.Path(self.qr_code_path.name).joinpath(token).with_suffix('.png')
            self.qr_generator.get_image(text=text, output_path=image_file_path)

            self._add_img(cell, str(image_file_path.absolute()), width, height, text)


        # Save the document
        document.save(output_path)
    
    def _add_img(self, cell, img_path, width, height, text):
            run = cell.paragraphs[0].add_run()
            run.add_picture(img_path, width= Mm(40), height= Mm(40))
            run.add_text(f'\n{text}')

    def _set_document_margins(self, document, top_margin, bottom_margin, left_margin, right_margin):
        sections = document.sections
        for section in sections:
            section.top_margin = Mm(top_margin)
            section.bottom_margin = Mm(bottom_margin)
            section.left_margin = Mm(left_margin)
            section.right_margin = Mm(right_margin)


if __name__ == '__main__':
    

    filepath = pathlib.Path('data/output/invitation_data.txt')
    output_path = 'qr_code_labels.docx'

    url = 'https://wedding.molitoris.org'

    with open(file=filepath, mode='r') as f:
        json_data = json.load(f)

    generator = QrCodeLabelGenerator()

    generator.create_document(json_data, url=url, output_path=output_path)